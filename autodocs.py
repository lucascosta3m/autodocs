# -*- coding: utf-8 -*-
import tkinter as tk
from tkinter import messagebox
import gspread
from google.oauth2.service_account import Credentials
from gspread.utils import rowcol_to_a1
from docx import Document
from pathlib import Path
import re
import os
import sys
import time
import openpyxl                     # Para manipular arquivos Excel .xlsx
from openpyxl.utils.exceptions import InvalidFileException # Para tratar erros de arquivo Excel inválido

# ==============================================================================
# 1. FUNÇÃO AUXILIAR PARA CAMINHOS (PyInstaller)
# ==============================================================================
def resource_path(relative_path):
    """ Obtém o caminho absoluto para o recurso, funciona para dev e PyInstaller """
    try:
        # PyInstaller cria uma pasta temporária e armazena o caminho em _MEIPASS
        base_path = Path(sys._MEIPASS)
    except Exception:
        # Se não estiver rodando via PyInstaller, pega o diretório do script
        try:
            base_path = Path(__file__).parent
        except NameError:
            # Fallback para o diretório atual se __file__ não estiver definido (ex: console interativo)
            base_path = Path(".")
    return base_path / relative_path

# ==============================================================================
# 2. CONFIGURAÇÕES GLOBAIS E CONSTANTES
# ==============================================================================
DEBUG_MODE = False # Mude para True para ver mais logs detalhados no console

# --- Caminhos e Arquivos ---
PASTA_SAIDA = "docs_gerados" # Será criada onde o script/exe rodar
PASTA_TEMPLATES_REL = "templates"
CAMINHO_CREDENCIAL_REL = "credenciais.json"

# Converte caminhos relativos para absolutos (funciona em dev e no .exe)
PASTA_TEMPLATES = resource_path(PASTA_TEMPLATES_REL)
CAMINHO_CREDENCIAL = resource_path(CAMINHO_CREDENCIAL_REL)

# --- Nomes Google Sheets ---
PLANILHA_PF_FILENAME = "FORMULÁRIO PESSOA FÍSICA (respostas)"
PLANILHA_PJ_FILENAME = "FORMULÁRIO PESSOA JURÍDICA (respostas)"
# Use os nomes exatos das abas (planilhas internas)
PF_TAB_NAME = "Respostas ao formulário PF"
PJ_TAB_NAME = "Respostas ao formulário PJ"

# --- Templates DOCX ---
TEMPLATE_PF_DOCX = [PASTA_TEMPLATES / f"PF{i}- {name}.docx" for i, name in enumerate(["FICHA DE INSCRICAO", "INSTRUMENTO", "PROPOSTA DE ADMISSAO", "TERMO DE ADESAO"], 1)]
TEMPLATE_PJ_DOCX = [PASTA_TEMPLATES / f"PJ{i}- {name}.docx" for i, name in enumerate(["FICHA DE INSCRICAO", "INSTRUMENTO", "PROPOSTA DE ADMISSAO", "TERMO DE ADESAO"], 1)]

# --- Templates XLSX ---
TEMPLATE_PF_XLSX = [PASTA_TEMPLATES / "PF - FORMULARIO DE MOBILIZAÇÃO.xlsx"]
TEMPLATE_PJ_XLSX = [PASTA_TEMPLATES / "PJ - FORMULARIO DE MOBILIZAÇÃO.xlsx"]

# --- Colunas Chave Planilha ---
TRIGGER_VALUE = "JÁ CADASTREI MEUS DADOS PESSOAIS, QUERO CADASTRAR OUTRO VEÍCULO"
COL_CADASTRO = "CADASTRO" # Coluna que indica se é um novo cadastro ou só veículo adicional
COL_PF_ID_TRIGGER = "CPF - (SOMENTE NÚMERO)" # Coluna com CPF/CNPJ na linha de veículo adicional
COL_PF_ID_COMPARISON = "CPF (somente número)" # Coluna com CPF/CNPJ principal (usado para buscar a fonte)
COL_PJ_ID_TRIGGER = "CNPJ - (SOMENTE NÚMERO)"
COL_PJ_ID_COMPARISON = "CNPJ (somente número)"
STATUS_COL = "Status" # Coluna para marcar como "GERADO"

# ==============================================================================
# 3. FUNÇÕES AUXILIARES (Formatação, Substituição)
# ==============================================================================
def formatar_cpf(cpf_input):
    """Formata CPF para xxx.xxx.xxx-xx."""
    if cpf_input is None: return ""
    cpf = str(cpf_input).strip().lstrip("'") # Remove apóstrofo inicial comum em planilhas
    cpf_limpo = re.sub(r'\D', '', cpf).zfill(11) # Remove não-dígitos e preenche com zeros à esquerda
    if len(cpf_limpo) == 11: return f'{cpf_limpo[:3]}.{cpf_limpo[3:6]}.{cpf_limpo[6:9]}-{cpf_limpo[9:]}'
    # if DEBUG_MODE: print(f"DBG: CPF inválido para formatação: '{cpf_input}'")
    return str(cpf_input) # Retorna original se inválido

def formatar_cnpj(cnpj_input):
    """Formata CNPJ para xx.xxx.xxx/xxxx-xx."""
    if cnpj_input is None: return ""
    cnpj = str(cnpj_input).strip().lstrip("'")
    cnpj_limpo = re.sub(r'\D', '', cnpj).zfill(14)
    if len(cnpj_limpo) == 14: return f'{cnpj_limpo[:2]}.{cnpj_limpo[2:5]}.{cnpj_limpo[5:8]}/{cnpj_limpo[8:12]}-{cnpj_limpo[12:]}'
    # if DEBUG_MODE: print(f"DBG: CNPJ inválido para formatação: '{cnpj_input}'")
    return str(cnpj_input)

def _criar_dicionario_placeholders(dados):
    """Cria o dicionário de placeholders {CHAVE_MAIUSCULA} -> valor formatado."""
    # Mapeia colunas específicas para suas funções de formatação
    formatadores = {
        COL_PF_ID_COMPARISON.upper(): formatar_cpf,
        COL_PJ_ID_COMPARISON.upper(): formatar_cnpj,
        COL_PF_ID_TRIGGER.upper(): formatar_cpf,
        COL_PJ_ID_TRIGGER.upper(): formatar_cnpj,
        # Adicione outras colunas que precisam de formatação específica aqui
        # Ex: 'TELEFONE'.upper(): formatar_telefone,
    }
    placeholders = {}
    for chave, valor in dados.items():
        # Ignora chaves internas adicionadas pelo script
        if str(chave).lower() in ["tipo", "linha"]:
            continue
        chave_fmt = str(chave).strip().upper() # Chave sempre maiúscula e sem espaços extras
        valor_fmt = str(valor).strip() if valor is not None else "" # Valor como string sem espaços extras

        formatador = formatadores.get(chave_fmt)
        if formatador:
            valor_fmt = formatador(valor_fmt) # Aplica formatação específica
        # Aplica UPPERCASE apenas se não foi formatado E não for puramente numérico
        elif not formatador and not re.fullmatch(r'\d+', valor_fmt):
             valor_fmt = valor_fmt.upper() # Converte para maiúsculas textos não formatados

        placeholders[f"{{{chave_fmt}}}"] = valor_fmt # Cria a chave no formato {NOME_COLUNA}
    return placeholders

def substituir_placeholders(document, dados):
    """Substitui placeholders {CHAVE} no documento DOCX, tentando preservar formatação."""
    placeholders = _criar_dicionario_placeholders(dados)
    if DEBUG_MODE: print("  -- Substituindo placeholders em DOCX --")

    # Função interna para processar 'runs' (partes de texto com a mesma formatação)
    def substituir_em_runs(runs):
        # Concatena o texto de todos os runs para encontrar placeholders que podem estar divididos
        full_text = "".join(run.text for run in runs)
        # Otimização: Se nenhum placeholder está presente no texto concatenado, não faz nada
        if not any(ph in full_text for ph in placeholders):
             return False # Indica que não houve modificação

        modified = False
        current_text = ""
        # Concatena o texto e limpa os runs originais
        for run in runs:
            current_text += run.text
            run.text = ""

        # Aplica as substituições no texto completo concatenado
        modified_text = current_text
        for ph, val in placeholders.items():
             # Usa replace para substituir todas as ocorrências do placeholder
             new_text = modified_text.replace(ph, str(val))
             if new_text != modified_text:
                 modified = True
                 modified_text = new_text

        # Reescreve o texto modificado no primeiro run da sequência
        # Isso geralmente preserva a formatação do primeiro run para todo o texto substituído
        if runs:
            runs[0].text = modified_text
        return modified

    # Itera sobre parágrafos no corpo principal do documento
    for paragraph in document.paragraphs:
        substituir_em_runs(paragraph.runs)

    # Itera sobre tabelas, células e parágrafos dentro das células
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    substituir_em_runs(paragraph.runs)


def substituir_placeholders_excel(workbook, dados):
    """Substitui placeholders {CHAVE} nas células do workbook Excel."""
    placeholders = _criar_dicionario_placeholders(dados)
    if DEBUG_MODE: print("  -- Substituindo placeholders em EXCEL --")
    modificado_geral = False # Flag para indicar se alguma célula foi alterada no workbook

    # Itera sobre todas as planilhas (abas) dentro do arquivo Excel
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        modificado_sheet = False
        # Itera sobre todas as linhas e células da planilha atual
        for row in sheet.iter_rows():
            for cell in row:
                # Verifica se a célula contém uma string (placeholders só funcionam em texto)
                if cell.value and isinstance(cell.value, str):
                    original_value = cell.value
                    modified_value = original_value
                    # Itera sobre os placeholders e substitui se encontrar na string da célula
                    for ph, val in placeholders.items():
                        if ph in modified_value:
                            # Substitui todas as ocorrências do placeholder
                            modified_value = modified_value.replace(ph, str(val))

                    # Se o valor foi modificado, atualiza a célula
                    if modified_value != original_value:
                        cell.value = modified_value
                        modificado_sheet = True # Marca que esta planilha foi modificada
        if modificado_sheet:
            modificado_geral = True
            if DEBUG_MODE: print(f"    Substituições realizadas na planilha '{sheet_name}'.")

    # if not modificado_geral and DEBUG_MODE:
    #      print("    Nenhuma substituição realizada no Excel.")
    return workbook # Retorna o workbook modificado


# ==============================================================================
# 4. FUNÇÃO DE PRÉ-PREENCHIMENTO
# ==============================================================================
def preencher_e_atualizar_planilha(sheet, headers, dados_originais, id_col_trigger, id_col_comparison, col_cadastro, trigger_value):
    """Preenche dados baseados em cadastros anteriores e atualiza planilha via API."""
    if not dados_originais:
        print(f"Aviso Pré-preenchimento: Planilha '{sheet.title}' vazia ou sem dados.")
        return []

    print(f"\n--- Iniciando pré-preenchimento para: {sheet.title} ---")
    try:
        # Valida se as colunas essenciais para o pré-preenchimento existem nos cabeçalhos lidos
        colunas_essenciais = [col_cadastro, id_col_trigger, id_col_comparison]
        for col in colunas_essenciais:
            if col not in headers: raise ValueError(f"Coluna essencial '{col}' não encontrada.")

        idx_col_comparison = headers.index(id_col_comparison) + 1 # Índice 1-based da coluna de comparação
        col_indices = {name: i + 1 for i, name in enumerate(headers)} # Mapeia nome -> índice 1-based
        # Define colunas que podem ser copiadas da linha fonte (ignora as de controle, status, etc.)
        cols_to_fill = { name: index for name, index in col_indices.items()
                         if name not in [col_cadastro, id_col_trigger, id_col_comparison, STATUS_COL, "Timestamp", "Carimbo de data/hora"] }
        if DEBUG_MODE: print(f" Colunas a preencher da fonte (se vazias na linha alvo): {list(cols_to_fill.keys())}")

    except ValueError as e:
        messagebox.showerror("Erro Config Pré-preenchimento", f"{e} na planilha '{sheet.title}'. Verifique nomes das colunas.")
        return dados_originais # Retorna dados originais se houve erro nos headers

    source_data_map = {} # Dicionário para guardar a primeira linha fonte encontrada para cada ID
    dados_modificados = [] # Lista para guardar os dados como lidos, mas potencialmente atualizados em memória
    linhas_para_buscar = [] # Lista de tuplas (índice_lista, trigger_id, linha_planilha) para linhas alvo

    print(" Mapeando dados de origem e identificando alvos...")
    for i, row_dict in enumerate(dados_originais):
        # Limpa chaves e valores (remove espaços, None vira "", ignora chaves vazias)
        clean_row = {str(k).strip(): str(v).strip() if v is not None else "" for k, v in row_dict.items() if k}
        dados_modificados.append(clean_row) # Adiciona o dict limpo
        row_num = i + 2 # Número da linha na planilha (1-based + cabeçalho)
        cadastro_status = clean_row.get(col_cadastro, "")
        comp_id = clean_row.get(id_col_comparison, "") # ID da linha (pode ser fonte)
        trig_id = clean_row.get(id_col_trigger, "")   # ID que a linha alvo usa para buscar a fonte

        # Se não for linha de trigger e tiver ID de comparação, é uma potencial fonte
        if cadastro_status != trigger_value and comp_id:
            if comp_id not in source_data_map: # Guarda apenas a primeira ocorrência como fonte
                 source_data_map[comp_id] = clean_row
        # Se for linha de trigger e tiver ID de trigger, é um alvo
        elif cadastro_status == trigger_value and trig_id:
            linhas_para_buscar.append((i, trig_id, row_num))
        elif cadastro_status == trigger_value and not trig_id:
            print(f" Aviso Pré-preenchimento: Linha {row_num} ({sheet.title}) com trigger ('{trigger_value}') mas sem ID em '{id_col_trigger}'. Não será pré-preenchida.")

    batch_updates = [] # Lista para guardar as atualizações a serem enviadas para a API do Google Sheets
    print(f" Processando {len(linhas_para_buscar)} linha(s) alvo para pré-preenchimento...")
    for list_idx, trigger_id, row_num in linhas_para_buscar:
        if list_idx >= len(dados_modificados): continue # Segurança
        target_row = dados_modificados[list_idx] # Pega o dicionário da linha alvo
        row_needs_api_update = False # Flag para saber se esta linha precisa de update via API

        # Passo 1: Preencher ID de Comparação na linha Alvo (se estiver vazio) com o ID Trigger
        current_comparison_val = target_row.get(id_col_comparison, "")
        if not current_comparison_val and trigger_id:
            target_row[id_col_comparison] = trigger_id # Atualiza em memória
            cell_a1 = rowcol_to_a1(row_num, idx_col_comparison) # Converte para A1 notation
            batch_updates.append({'range': cell_a1, 'values': [[str(trigger_id)]]}) # Prepara API update
            row_needs_api_update = True
            if DEBUG_MODE: print(f"  L{row_num}: Preenchendo '{id_col_comparison}' com ID Trigger '{trigger_id}'")

        # Passo 2: Preencher outros campos da linha Alvo (se vazios) buscando da linha Fonte
        if trigger_id in source_data_map:
            source_row = source_data_map[trigger_id]
            if DEBUG_MODE and not row_needs_api_update: print(f"  L{row_num}: Fonte encontrada para ID '{trigger_id}'. Verificando campos...")
            campos_preenchidos_da_fonte = 0
            for col_name, col_idx in cols_to_fill.items():
                # Verifica se o campo está VAZIO na linha Alvo
                if not target_row.get(col_name, ""):
                    source_value = source_row.get(col_name, "")
                    # Verifica se o campo tem valor na linha Fonte
                    if source_value:
                        target_row[col_name] = source_value # Atualiza em memória
                        cell_a1 = rowcol_to_a1(row_num, col_idx)
                        batch_updates.append({'range': cell_a1, 'values': [[str(source_value)]]}) # Prepara API update
                        row_needs_api_update = True
                        campos_preenchidos_da_fonte += 1
            if DEBUG_MODE and campos_preenchidos_da_fonte > 0:
                 print(f"    L{row_num}: Preenchidos {campos_preenchidos_da_fonte} campos a partir da fonte.")

        elif DEBUG_MODE and not row_needs_api_update: # Se não achou fonte E não preencheu o ID no passo 1
             print(f"  L{row_num}: Fonte não encontrada para ID '{trigger_id}' e coluna de comparação já estava preenchida. Nenhum pré-preenchimento realizado.")

    # Envia as atualizações em lote para a API, se houver alguma
    if batch_updates:
        print(f" Enviando {len(batch_updates)} atualizações de pré-preenchimento para '{sheet.title}'...")
        try:
            sheet.batch_update(batch_updates, value_input_option='USER_ENTERED')
            print(" Pré-preenchimento salvo com sucesso na planilha!")
        except gspread.exceptions.APIError as e:
            print(f"ERRO de API ao salvar pré-preenchimento em '{sheet.title}': {e}")
            messagebox.showerror("Erro API Google (Pré-preenchimento)", f"Falha ao salvar pré-preenchimento em '{sheet.title}':\n{e}")
        except Exception as e:
             print(f"ERRO inesperado ao salvar pré-preenchimento em '{sheet.title}': {e}")
             messagebox.showerror("Erro Inesperado (Pré-preenchimento)", f"Falha ao salvar pré-preenchimento em '{sheet.title}':\n{e}")
    else:
        print(f" Nenhuma atualização de pré-preenchimento necessária para enviar à API para '{sheet.title}'.")

    print(f"--- Fim do pré-preenchimento para: {sheet.title} ---\n")
    return dados_modificados # Retorna a lista de dicionários potencialmente modificada em memória

# ==============================================================================
# 5. FUNÇÕES DE INTERAÇÃO COM PLANILHAS (Carregar)
# ==============================================================================
def autenticar_google():
    """Autentica com a API do Google Sheets e retorna o cliente gspread."""
    print("Autenticando com Google...")
    scopes = ['https://www.googleapis.com/auth/spreadsheets', 'https://www.googleapis.com/auth/drive']
    if not CAMINHO_CREDENCIAL.exists():
         raise FileNotFoundError(f"Arquivo de credenciais não encontrado em: {CAMINHO_CREDENCIAL}")
    credenciais = Credentials.from_service_account_file(CAMINHO_CREDENCIAL, scopes=scopes)
    gc = gspread.authorize(credenciais)
    print("Autenticado.")
    return gc

def carregar_planilha(gc, filename, tab_name):
    """Abre uma planilha e aba específica, lê cabeçalho e dados de forma robusta."""
    print(f"\nAbrindo arquivo: {filename}")
    try:
        workbook = gc.open(filename)
        print(f"Tentando acessar aba: '{tab_name}'")
        sheet = workbook.worksheet(tab_name)
        print(f"Aba '{sheet.title}' (ID: {sheet.id}) acessada.")

        # Lê a primeira linha (cabeçalho) BRUTA
        raw_headers_list = sheet.row_values(1)
        if not raw_headers_list:
            raise ValueError("Cabeçalho (Linha 1) está completamente vazio ou não pôde ser lido.")

        # Limpa os cabeçalhos (remove espaços extras das pontas, converte para string)
        cleaned_headers_list = [str(h).strip() for h in raw_headers_list]
        # Filtra cabeçalhos que ficaram vazios APÓS limpeza
        headers_non_empty = [h for h in cleaned_headers_list if h]
        if not headers_non_empty:
             raise ValueError("Todos os cabeçalhos na Linha 1 estão vazios após limpeza.")
        if DEBUG_MODE:
             print(f"Cabeçalho final (limpo, não-vazio) a ser usado: {headers_non_empty}")

        # Verifica Duplicatas nos cabeçalhos NÃO VAZIOS (case-insensitive)
        headers_lower_non_empty = [h.lower() for h in headers_non_empty]
        if len(headers_lower_non_empty) != len(set(headers_lower_non_empty)):
            from collections import Counter
            header_counts = Counter(headers_lower_non_empty)
            duplicates = {header: count for header, count in header_counts.items() if count > 1}
            # Levanta um erro claro se duplicatas forem encontradas aqui
            raise gspread.exceptions.GSpreadException(
                f"Cabeçalho não único detectado programaticamente na aba '{sheet.title}'. "
                f"Duplicatas (minúsculas): {duplicates}. Corrija a Linha 1 da planilha."
            )

        # Valida se a coluna Status Essencial está presente nos cabeçalhos finais
        if STATUS_COL not in headers_non_empty:
            raise ValueError(f"Coluna Status Essencial '{STATUS_COL}' não encontrada no cabeçalho final (após limpeza e remoção de vazios).")

        print(f"Lendo dados da aba '{sheet.title}' usando get_all_records...")
        # Usa get_all_records passando os cabeçalhos limpos e validados como 'expected_headers'
        # Isso ajuda gspread a mapear corretamente e evita erros de 'cabeçalho não único' causados por problemas de detecção interna
        data = sheet.get_all_records(head=1, expected_headers=headers_non_empty)
        print(f"Lidos {len(data)} registros de '{sheet.title}' com sucesso.")

        return sheet, headers_non_empty, data # Retorna os headers limpos e não-vazios

    except gspread.exceptions.SpreadsheetNotFound:
        messagebox.showerror("Erro Crítico", f"Arquivo '{filename}' não encontrado no Google Drive. Verifique o nome e as permissões da conta de serviço.")
        sys.exit()
    except gspread.exceptions.WorksheetNotFound:
        messagebox.showerror("Erro Crítico", f"Aba '{tab_name}' não encontrada no arquivo '{filename}'. Verifique o nome exato (sensível a maiúsculas/minúsculas e espaços).")
        sys.exit()
    except ValueError as e: # Erro levantado pela nossa validação de cabeçalho/Status
         messagebox.showerror("Erro Cabeçalho/Validação", f"Erro no cabeçalho/validação da aba '{tab_name}' em '{filename}':\n{e}")
         sys.exit()
    except gspread.exceptions.GSpreadException as ge: # Erro de duplicata ou outro erro gspread
         print(f"ERRO GSPREAD FINAL: {ge}") # Loga o erro
         messagebox.showerror("Erro GSpread", f"Erro ao processar a planilha '{filename}' / Aba '{tab_name}':\n{ge}\n\nVerifique a formatação da planilha ou o console para mais detalhes.")
         sys.exit()
    except Exception as e: # Outros erros (API, rede, etc.)
        messagebox.showerror("Erro ao Carregar Planilha", f"Erro inesperado ao carregar '{filename}' / Aba '{tab_name}':\n{type(e).__name__}: {e}")
        sys.exit()


# ==============================================================================
# 6. FUNÇÕES PRINCIPAIS DA APLICAÇÃO (Gerar Docs, Excluir)
# ==============================================================================

# Variáveis globais para acesso pelas funções de comando dos botões
# (Alternativa seria usar uma classe App, mas mantendo funcional por ora)
root = None
checkboxes_pf = []
checkboxes_pj = []
dados_pf_processados = []
dados_pj_processados = []
sheet_pf = None
sheet_pj = None
col_index_status_pf = -1
col_index_status_pj = -1

def gerar_documentos_cmd():
    """Função chamada pelo botão 'Gerar Documentos'. Processa DOCX e XLSX."""
    global checkboxes_pf, checkboxes_pj, sheet_pf, sheet_pj, col_index_status_pf, col_index_status_pj, root

    selecionados_tuplas = [(p, var) for p, var in checkboxes_pf + checkboxes_pj if var.get()]
    if not selecionados_tuplas:
        messagebox.showwarning("Aviso", "Nenhum registro selecionado.")
        return

    print(f"\n--- Gerando Docs para {len(selecionados_tuplas)} registro(s) ---")
    try:
        # Cria a pasta de saída se não existir
        Path(PASTA_SAIDA).mkdir(parents=True, exist_ok=True)
    except OSError as e:
        messagebox.showerror("Erro ao Criar Pasta", f"Não foi possível criar a pasta de saída '{PASTA_SAIDA}': {e}")
        return

    erros_template = [] # Guarda erros de templates não encontrados
    erros_geracao = []  # Guarda erros durante a geração/salvamento de arquivos individuais
    erros_atualizacao = [] # Guarda erros ao preparar/enviar atualização de status
    updates_status_batch = [] # Lista para guardar atualizações de status para a API

    # Itera sobre cada pessoa/empresa selecionada na interface
    for pessoa_original, var_tk in selecionados_tuplas:
        # Garante que as chaves do dicionário sejam strings e estejam em maiúsculo
        dados_pessoa = {str(k).strip().upper(): v for k, v in pessoa_original.items() if k}

        tipo = dados_pessoa.get('TIPO') # 'TIPO' foi adicionado no processamento
        linha = pessoa_original.get("linha") # 'linha' foi adicionada no processamento

        # Define quais listas de templates usar e qual planilha/coluna de status atualizar
        if tipo == "PF":
            templates_docx = TEMPLATE_PF_DOCX
            templates_xlsx = TEMPLATE_PF_XLSX
            ws, col_idx = sheet_pf, col_index_status_pf
        elif tipo == "PJ":
            templates_docx = TEMPLATE_PJ_DOCX
            templates_xlsx = TEMPLATE_PJ_XLSX
            ws, col_idx = sheet_pj, col_index_status_pj
        else:
            msg_erro = f"L{linha}: Tipo de pessoa ('{tipo}') inválido ou não encontrado nos dados."
            print(f"⚠️ {msg_erro}"); erros_geracao.append(msg_erro); continue # Pula este registro

        # Define nomes base para pastas e arquivos, tratando caracteres inválidos
        nome_base_raw = str(dados_pessoa.get("NOME COMPLETO", dados_pessoa.get("RAZÃO SOCIAL", f"Registro_L{linha}"))).strip()
        placa_base_raw = str(dados_pessoa.get("PLACA", "SemPlaca")).strip()
        # Remove caracteres inválidos para nomes de arquivo/pasta e limita comprimento
        nome_base = re.sub(r'[\\/*?:"<>|]', "", nome_base_raw).replace(" ", "_")[:80]
        placa_base = re.sub(r'[\\/*?:"<>|]', "", placa_base_raw).replace('-', '')

        if not isinstance(linha, int) or linha < 2:
            msg_erro = f"{nome_base}: Número de linha inválido ({linha}) associado ao registro."
            print(f"⚠️ {msg_erro}"); erros_geracao.append(msg_erro); continue

        if DEBUG_MODE: print(f"\n=== PROC REG: {nome_base_raw} (L{linha}, {tipo}) ===")
        pasta_destino = Path(PASTA_SAIDA) / nome_base
        try:
             pasta_destino.mkdir(parents=True, exist_ok=True)
        except OSError as e:
             msg = f"L{linha} ({nome_base_raw}): Erro ao criar pasta de destino '{pasta_destino}': {e}"; print(f"❌ {msg}"); erros_geracao.append(msg); continue

        # Flag para controlar se TODOS os templates (docx e xlsx) foram gerados com sucesso para esta pessoa
        todos_templates_ok_para_pessoa = True

        # --- Processamento dos Templates DOCX ---
        if DEBUG_MODE and templates_docx: print(f"  -- Processando {len(templates_docx)} templates DOCX --")
        for template_path_obj in templates_docx:
            template_nome = template_path_obj.name
            if not template_path_obj.exists():
                msg = f"Template DOCX não encontrado: {template_nome}"; print(f"⚠️ {msg}"); erros_template.append(msg); todos_templates_ok_para_pessoa = False; continue
            if DEBUG_MODE: print(f"-- Proc Template DOCX: {template_nome}")
            try:
                doc = Document(template_path_obj)
                substituir_placeholders(doc, dados_pessoa)
                prefixo = template_path_obj.stem.split('-', 1)[0].strip()
                nome_doc = f"{prefixo}_{nome_base}_{placa_base}.docx"
                caminho_saida = pasta_destino / nome_doc
                doc.save(caminho_saida)
                # if DEBUG_MODE: print(f"  >> Salvo DOCX: {nome_doc}")
            except Exception as e:
                msg = f"L{linha} ({nome_base_raw}): Erro ao gerar/salvar DOCX '{template_nome}': {type(e).__name__} - {e}"; print(f"❌ {msg}"); erros_geracao.append(msg); todos_templates_ok_para_pessoa = False

        # --- Processamento dos Templates XLSX ---
        if DEBUG_MODE and templates_xlsx: print(f"  -- Processando {len(templates_xlsx)} templates XLSX --")
        for template_path_obj in templates_xlsx:
            template_nome = template_path_obj.name
            if not template_path_obj.exists():
                 msg = f"Template XLSX não encontrado: {template_nome}"; print(f"⚠️ {msg}"); erros_template.append(msg); todos_templates_ok_para_pessoa = False; continue
            if DEBUG_MODE: print(f"-- Proc Template XLSX: {template_nome}")
            try:
                workbook = openpyxl.load_workbook(template_path_obj)
                workbook_modificado = substituir_placeholders_excel(workbook, dados_pessoa)
                prefixo_xlsx = template_path_obj.stem # Nome do arquivo template sem extensão
                nome_doc_xlsx = f"{prefixo_xlsx}_{nome_base}_{placa_base}.xlsx"
                caminho_saida_xlsx = pasta_destino / nome_doc_xlsx
                workbook_modificado.save(caminho_saida_xlsx)
                # if DEBUG_MODE: print(f"  >> Salvo XLSX: {nome_doc_xlsx}")
            except InvalidFileException:
                 msg = f"L{linha} ({nome_base_raw}): Arquivo Excel inválido ou corrompido: '{template_nome}'"; print(f"❌ {msg}"); erros_geracao.append(msg); todos_templates_ok_para_pessoa = False
            except Exception as e:
                 msg = f"L{linha} ({nome_base_raw}): Erro ao gerar/salvar XLSX '{template_nome}': {type(e).__name__} - {e}"; print(f"❌ {msg}"); erros_geracao.append(msg); todos_templates_ok_para_pessoa = False

        # --- Preparar Atualização de Status (APENAS se TUDO deu certo para esta pessoa) ---
        if todos_templates_ok_para_pessoa:
            if col_idx > 0: # Verifica se o índice da coluna Status é válido
                if DEBUG_MODE: print(f"  >> TODOS Docs OK ({nome_base_raw}). Preparando update status '{ws.title}' L{linha} C{col_idx}")
                try:
                    cell_a1 = rowcol_to_a1(linha, col_idx)
                    updates_status_batch.append({'range': cell_a1, 'values': [['GERADO']], 'worksheet': ws})
                except Exception as e:
                    msg = f"L{linha} ({nome_base_raw}): Erro ao preparar A1 para status ({linha},{col_idx}): {e}"; print(f"❌ {msg}"); erros_atualizacao.append(msg)
            else:
                msg = f"L{linha} ({nome_base_raw}): Docs gerados, mas coluna '{STATUS_COL}' não encontrada ou inválida em '{ws.title}'. Status não será atualizado."; print(f"⚠️ {msg}"); erros_atualizacao.append(msg)
        elif DEBUG_MODE:
            print(f"  >> Geração INCOMPLETA/ERRO para {nome_base_raw}. Status NÃO será atualizado.")


    # --- Envio das Atualizações de Status em Lote ---
    status_atualizado_count = 0
    if updates_status_batch:
        print(f"\n--- Enviando {len(updates_status_batch)} updates de status para Google Sheets ---")
        updates_agrupados = {} # Agrupa por objeto worksheet
        for info in updates_status_batch:
            ws = info['worksheet']
            if ws not in updates_agrupados: updates_agrupados[ws] = []
            updates_agrupados[ws].append({'range': info['range'], 'values': info['values']})

        for sheet_obj, updates_list in updates_agrupados.items():
            title = sheet_obj.title
            print(f" Enviando {len(updates_list)} updates para '{title}'...")
            try:
                response = sheet_obj.batch_update(updates_list, value_input_option='USER_ENTERED')
                print(f"   -> Sucesso ao enviar para '{title}'.")
                status_atualizado_count += len(updates_list)
            except gspread.exceptions.APIError as e:
                 msg = f"   -> ERRO API Google ao atualizar status em '{title}': {e}"; print(msg); erros_atualizacao.append(msg)
            except Exception as e:
                 msg = f"   -> ERRO inesperado ao atualizar status em '{title}': {e}"; print(msg); erros_atualizacao.append(msg)
        print("--- Fim do envio de status ---")
    else:
        print("\nNenhum update de status a enviar.")

    # --- Monta Relatório Final para o Usuário ---
    msg_final = [f"Processo Concluído.", f"Registros Selecionados: {len(selecionados_tuplas)}", f"Status 'GERADO' atualizado (API OK): {status_atualizado_count}"]
    if erros_template: msg_final.extend(["\n--- Templates Não Encontrados ---"] + list(set(erros_template)))
    if erros_geracao: msg_final.extend([f"\n--- Erros Geração/Salvar Docs ({len(erros_geracao)}) ---"] + erros_geracao[:5] + ["(... ver console para mais detalhes)"] if len(erros_geracao) > 5 else erros_geracao)
    if erros_atualizacao: msg_final.extend([f"\n--- Erros Preparação/Envio Status ({len(erros_atualizacao)}) ---"] + erros_atualizacao[:5] + ["(... ver console para mais detalhes)"] if len(erros_atualizacao) > 5 else erros_atualizacao)
    messagebox.showinfo("Relatório Final da Geração", "\n".join(msg_final))

    # Fecha a janela atual para forçar recarga dos dados na próxima execução
    print("Recarregando interface...")
    try:
        if root and root.winfo_exists(): root.destroy()
    except tk.TclError: pass # Ignora erro se a janela já foi destruída


def excluir_entradas_cmd():
    """Função chamada pelo botão 'Excluir da Planilha'."""
    global checkboxes_pf, checkboxes_pj, sheet_pf, sheet_pj, root

    selecionados_tuplas = [(p, var) for p, var in checkboxes_pf + checkboxes_pj if var.get()]
    if not selecionados_tuplas: messagebox.showwarning("Aviso", "Nenhum registro selecionado."); return
    confirm = messagebox.askyesno("Confirmar Exclusão", f"Tem certeza que deseja excluir {len(selecionados_tuplas)} registro(s) da(s) planilha(s)?\n\nESTA AÇÃO NÃO PODE SER DESFEITA.")
    if not confirm: return

    print(f"\n--- Excluindo {len(selecionados_tuplas)} registro(s) ---")
    excluidos_count = 0; erros_exclusao = []; requests_pf = []; requests_pj = []
    # Ordena por linha DECRESCENTE para evitar problemas de índice na exclusão em lote
    selecionados_ordenados = sorted(selecionados_tuplas, key=lambda item: item[0].get("linha", 0), reverse=True)
    try:
        sheet_id_pf = sheet_pf.id; sheet_id_pj = sheet_pj.id # IDs internos das abas
    except Exception as e:
        messagebox.showerror("Erro Interno", f"Erro ao obter IDs das abas para exclusão: {e}"); return

    # Prepara as requisições de exclusão para a API batchUpdate
    for pessoa_dict, _ in selecionados_ordenados:
        linha = pessoa_dict.get("linha"); tipo = pessoa_dict.get("tipo")
        nome = str(pessoa_dict.get("NOME COMPLETO", pessoa_dict.get("RAZÃO SOCIAL", "Reg Desconhecido"))).strip()
        if not isinstance(linha, int) or linha < 2:
             msg = f"Exclusão: Linha inválida ({linha}) para {nome}."; print(f"⚠️ {msg}"); erros_exclusao.append(msg); continue

        start_idx = linha - 1; end_idx = linha # API usa índice 0-based
        delete_req = {'deleteDimension': {'range': {'sheetId': None, 'dimension': 'ROWS', 'startIndex': start_idx, 'endIndex': end_idx}}}
        if tipo == "PF":
             delete_req['deleteDimension']['range']['sheetId'] = sheet_id_pf; requests_pf.append(delete_req)
        elif tipo == "PJ":
             delete_req['deleteDimension']['range']['sheetId'] = sheet_id_pj; requests_pj.append(delete_req)
        else:
             msg = f"Exclusão: Tipo '{tipo}' desconhecido L{linha}."; print(f"⚠️ {msg}"); erros_exclusao.append(msg)

    # Função auxiliar para executar o batchUpdate de exclusão por planilha
    def executar_exclusao(sheet_obj, requests, tipo_label):
        count = 0
        if requests:
            print(f" Executando {len(requests)} exclusões na planilha {tipo_label} ('{sheet_obj.title}')...")
            try:
                # Exclusão é feita no objeto spreadsheet, não no worksheet
                sheet_obj.spreadsheet.batch_update({'requests': requests})
                print(f"   > Exclusões {tipo_label} concluídas."); count = len(requests)
            except gspread.exceptions.APIError as e:
                 msg = f"Erro API Google ao excluir {tipo_label}: {e}"; print(f"❌ {msg}"); erros_exclusao.append(msg)
            except Exception as e:
                 msg = f"Erro inesperado ao excluir {tipo_label}: {e}"; print(f"❌ {msg}"); erros_exclusao.append(msg)
        return count

    # Executa as exclusões para PF e PJ
    excluidos_count += executar_exclusao(sheet_pf, requests_pf, "PF")
    excluidos_count += executar_exclusao(sheet_pj, requests_pj, "PJ")

    # Mostra relatório final da exclusão
    msg_final = [f"{excluidos_count} registro(s) efetivamente excluído(s) (comando API enviado)."]
    if erros_exclusao: msg_final.extend(["\n--- Ocorrências Durante Exclusão ---"] + erros_exclusao)
    messagebox.showinfo("Relatório de Exclusão", "\n".join(msg_final))

    # Fecha a janela para recarregar
    print("Recarregando interface após exclusão...")
    try:
        if root and root.winfo_exists(): root.destroy()
    except tk.TclError: pass


# ==============================================================================
# 7. FUNÇÕES DA INTERFACE GRÁFICA (Tkinter)
# ==============================================================================
def adicionar_checkbox(pessoa_data, parent_frame, checkboxes_list):
    """Cria e adiciona um checkbox para uma pessoa/empresa na interface."""
    var = tk.BooleanVar() # Variável Tkinter para controlar o estado (marcado/desmarcado)
    nome = str(pessoa_data.get("NOME COMPLETO", pessoa_data.get("RAZÃO SOCIAL", "N/A"))).strip()
    placa = str(pessoa_data.get("PLACA", "N/A")).strip()
    max_len_nome = 35 # Limita o tamanho do nome exibido para não alargar muito
    nome_display = f"{nome[:max_len_nome]}..." if len(nome) > max_len_nome else nome
    texto = f"{nome_display} - {placa}" # Texto do checkbox
    # Cria o checkbox dentro do frame pai (frame_pf ou frame_pj)
    cb = tk.Checkbutton(parent_frame, text=texto, variable=var, anchor="w", justify="left", wraplength=350)
    cb.pack(fill="x", padx=5, pady=1) # Adiciona ao frame pai
    checkboxes_list.append((pessoa_data, var)) # Guarda o dicionário de dados e a variável Tkinter

def criar_interface(root_window, dados_pf, dados_pj):
    """Cria todos os elementos da interface gráfica principal."""
    global checkboxes_pf, checkboxes_pj # Permite que esta função popule as listas globais

    # Configurações da janela principal
    root_window.title("Autodocs - Gerador de Documentos v1.3")
    root_window.geometry("850x650")
    root_window.minsize(700, 500)

    # --- Frame Principal ---
    main_frame = tk.Frame(root_window)
    main_frame.pack(fill=tk.BOTH, expand=True)

    # --- Botões Principais (Topo) ---
    frame_botoes_principais = tk.Frame(main_frame)
    frame_botoes_principais.pack(pady=10, fill="x")
    btn_gerar = tk.Button(frame_botoes_principais, text="Gerar Documentos Selecionados", command=gerar_documentos_cmd, width=30, height=2, bg="#D0F0D0", font=('Segoe UI', 10, 'bold')) # Cor verde clara
    btn_gerar.pack(side=tk.TOP, pady=5)
    # Botão de excluir foi removido/comentado em versões anteriores, mantendo assim.

    # --- Área Rolável (Centro) ---
    frame_scroll_container = tk.Frame(main_frame)
    frame_scroll_container.pack(pady=5, padx=10, fill="both", expand=True)
    canvas = tk.Canvas(frame_scroll_container)
    scrollbar_y = tk.Scrollbar(frame_scroll_container, orient=tk.VERTICAL, command=canvas.yview)
    scrollbar_x = tk.Scrollbar(frame_scroll_container, orient=tk.HORIZONTAL, command=canvas.xview)
    canvas.configure(yscrollcommand=scrollbar_y.set, xscrollcommand=scrollbar_x.set)
    scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
    scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
    canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    # Frame interno que conterá os checkboxes e rolará com o canvas
    frame_content = tk.Frame(canvas)
    canvas.create_window((0, 0), window=frame_content, anchor="nw", tags="frame_content")
    # Função para atualizar a região de rolagem quando o tamanho do frame interno mudar
    def on_frame_configure(event): canvas.configure(scrollregion=canvas.bbox("all"))
    frame_content.bind("<Configure>", on_frame_configure)

    # --- Frames PF e PJ (Dentro da Área Rolável, usando Grid) ---
    frame_content.columnconfigure(0, weight=1) # Coluna PF expande
    frame_content.columnconfigure(1, weight=1) # Coluna PJ expande
    frame_pf = tk.LabelFrame(frame_content, text=f"PESSOA FÍSICA ({len(dados_pf)})", padx=10, pady=10, font=('Segoe UI', 9, 'bold'))
    frame_pf.grid(row=0, column=0, padx=10, pady=5, sticky="nsew") # Ocupa espaço disponível
    frame_pj = tk.LabelFrame(frame_content, text=f"PESSOA JURÍDICA ({len(dados_pj)})", padx=10, pady=10, font=('Segoe UI', 9, 'bold'))
    frame_pj.grid(row=0, column=1, padx=10, pady=5, sticky="nsew")

    # --- Populando Checkboxes ---
    checkboxes_pf.clear(); checkboxes_pj.clear() # Limpa listas caso a interface seja recarregada
    print(f"\nPopulando Interface...")
    count_pf_visivel, count_pf_gerado = 0, 0
    for pessoa in dados_pf:
        if str(pessoa.get(STATUS_COL, "")).strip().upper() != "GERADO":
            adicionar_checkbox(pessoa, frame_pf, checkboxes_pf)
            count_pf_visivel += 1
        else: count_pf_gerado += 1

    count_pj_visivel, count_pj_gerado = 0, 0
    for pessoa in dados_pj:
        if str(pessoa.get(STATUS_COL, "")).strip().upper() != "GERADO":
            adicionar_checkbox(pessoa, frame_pj, checkboxes_pj)
            count_pj_visivel += 1
        else: count_pj_gerado += 1

    print(f" PF: {count_pf_visivel} para exibir ({count_pf_gerado} já 'GERADO'). Total: {len(dados_pf)}")
    print(f" PJ: {count_pj_visivel} para exibir ({count_pj_gerado} já 'GERADO'). Total: {len(dados_pj)}")
    # Atualiza títulos dos frames com contagem visível
    frame_pf.config(text=f"PESSOA FÍSICA ({count_pf_visivel} visíveis)")
    frame_pj.config(text=f"PESSOA JURÍDICA ({count_pj_visivel} visíveis)")


    # --- Botões de Seleção (Fundo, acima do label dev) ---
    frame_botoes_selecao = tk.Frame(main_frame)
    frame_botoes_selecao.pack(pady=(5,5), fill="x", side=tk.BOTTOM)
    def selecionar(lista): [var.set(True) for _, var in lista]
    def desmarcar(lista): [var.set(False) for _, var in lista]
    # Organiza botões em frames para melhor distribuição
    frame_sel_pf = tk.Frame(frame_botoes_selecao)
    frame_sel_pf.pack(side=tk.LEFT, expand=True, fill='x', padx=10)
    tk.Button(frame_sel_pf, text="Selecionar Todos PF", command=lambda: selecionar(checkboxes_pf), width=18).pack(side=tk.LEFT, padx=5)
    tk.Button(frame_sel_pf, text="Desmarcar Todos PF", command=lambda: desmarcar(checkboxes_pf), width=18).pack(side=tk.LEFT, padx=5)

    frame_sel_pj = tk.Frame(frame_botoes_selecao)
    frame_sel_pj.pack(side=tk.RIGHT, expand=True, fill='x', padx=10)
    tk.Button(frame_sel_pj, text="Selecionar Todos PJ", command=lambda: selecionar(checkboxes_pj), width=18).pack(side=tk.LEFT, padx=5)
    tk.Button(frame_sel_pj, text="Desmarcar Todos PJ", command=lambda: desmarcar(checkboxes_pj), width=18).pack(side=tk.LEFT, padx=5)

    # --- Label (Fundo) ---
    label_dev = tk.Label(main_frame, text='Desenvolvido por Lucas Costa', font=('Segoe UI', 8))
    label_dev.pack(pady=(0,5), side=tk.BOTTOM)


# ==============================================================================
# 8. EXECUÇÃO PRINCIPAL
# ==============================================================================
if __name__ == "__main__":
    try:
        # --- 1. Autenticação e Carregamento Inicial ---
        client_gspread = autenticar_google()
        sheet_pf, headers_pf, dados_pf_originais = carregar_planilha(client_gspread, PLANILHA_PF_FILENAME, PF_TAB_NAME)
        sheet_pj, headers_pj, dados_pj_originais = carregar_planilha(client_gspread, PLANILHA_PJ_FILENAME, PJ_TAB_NAME)

        # --- 2. Calcula Índices da Coluna de Status (1-based) ---
        try: col_index_status_pf = headers_pf.index(STATUS_COL) + 1
        except ValueError: messagebox.showerror("Erro Fatal", f"Coluna Status '{STATUS_COL}' não encontrada nos cabeçalhos da planilha PF."); sys.exit()
        try: col_index_status_pj = headers_pj.index(STATUS_COL) + 1
        except ValueError: messagebox.showerror("Erro Fatal", f"Coluna Status '{STATUS_COL}' não encontrada nos cabeçalhos da planilha PJ."); sys.exit()

        # --- 3. Pré-preenchimento ---
        dados_pf_preenchidos = preencher_e_atualizar_planilha(
            sheet_pf, headers_pf, dados_pf_originais,
            COL_PF_ID_TRIGGER, COL_PF_ID_COMPARISON, COL_CADASTRO, TRIGGER_VALUE
        )
        dados_pj_preenchidos = preencher_e_atualizar_planilha(
            sheet_pj, headers_pj, dados_pj_originais,
            COL_PJ_ID_TRIGGER, COL_PJ_ID_COMPARISON, COL_CADASTRO, TRIGGER_VALUE
        )

        # --- 4. Processamento Final dos Dados ---
        # Adiciona 'tipo' e 'linha' a cada dicionário para uso na interface e geração
        dados_pf_processados = [dict(row, tipo="PF", linha=i+2) for i, row in enumerate(dados_pf_preenchidos)]
        dados_pj_processados = [dict(row, tipo="PJ", linha=i+2) for i, row in enumerate(dados_pj_preenchidos)]

        # --- 5. Interface Gráfica ---
        root = tk.Tk()
        criar_interface(root, dados_pf_processados, dados_pj_processados)
        print("\nInterface pronta. Aguardando interação do usuário...")
        root.mainloop() # Mantém a janela aberta

    except FileNotFoundError as fnf_error:
         print(f"ERRO FATAL: Arquivo não encontrado - {fnf_error}")
         messagebox.showerror("Erro Fatal - Arquivo Não Encontrado", f"Não foi possível encontrar um arquivo essencial:\n{fnf_error}\n\nVerifique se o arquivo '{CAMINHO_CREDENCIAL_REL}' está na pasta correta.")
         sys.exit(1)
    except Exception as main_error:
        print(f"ERRO FATAL NA EXECUÇÃO PRINCIPAL: {type(main_error).__name__} - {main_error}")
        import traceback
        traceback.print_exc() # Imprime detalhes do erro no console
        try:
            messagebox.showerror("Erro Fatal Inesperado", f"Ocorreu um erro crítico:\n{type(main_error).__name__}: {main_error}\n\nVerifique o console para detalhes técnicos.")
        except Exception: pass # Ignora erro ao mostrar messagebox se o Tkinter falhou
        sys.exit(1)

    print("Aplicação finalizada.")